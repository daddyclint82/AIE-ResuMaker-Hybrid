from fastapi import FastAPI, Request, Form
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import json
import os
import httpx
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import base64
from io import BytesIO

from typing import Dict, List, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib import colors

from dotenv import load_dotenv

# Environment detection
APP_ENV = os.getenv("APP_ENV", "development").lower()
IS_PRODUCTION = APP_ENV == "production"

# Load the appropriate .env file
if IS_PRODUCTION:
    env_file = ".env.production"
else:
    # Development: prefer .env.development, fallback to .env
    env_file = ".env.development" if os.path.exists(".env.development") else ".env"

env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), env_file)

if os.path.exists(env_path):
    load_dotenv(env_path)
else:
    load_dotenv()

import stripe

# Stripe Configuration — loaded from environment
STRIPE_PUBLISHABLE_KEY = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_PRICE_ID = os.getenv("STRIPE_PRICE_ID", "")
STRIPE_STUDENT_PRICE_ID = os.getenv("STRIPE_STUDENT_PRICE_ID", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")

stripe.api_key = STRIPE_SECRET_KEY

# Base URL for redirects — must be set in production
BASE_URL = os.getenv("BASE_URL", "http://localhost:8000")

# Email Configuration
SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = os.getenv("SMTP_FROM", SMTP_USER)

def send_verification_email(to_email: str, token: str, base_url: str) -> bool:
    """Send student verification email with unique link"""
    try:
        if not SMTP_USER or not SMTP_PASS:
            print(f"[EMAIL] SMTP not configured. Token for {to_email}: {token}")
            return False
        
        # Create message
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'Verify Your Student Status - AIE ResuMaker'
        msg['From'] = SMTP_FROM
        msg['To'] = to_email
        
        # Verification link
        verify_link = f"{base_url}/verify-student?token={token}"
        
        # HTML email body
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <div style="background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%); padding: 30px; text-align: center; border-radius: 12px 12px 0 0;">
                <h1 style="color: white; margin: 0; font-size: 24px;">🎓 AIE ResuMaker</h1>
                <p style="color: #e0e7ff; margin: 10px 0 0 0;">Student Verification</p>
            </div>
            
            <div style="background: #f8fafc; padding: 30px; border-radius: 0 0 12px 12px; border: 1px solid #e2e8f0; border-top: none;">
                <p style="font-size: 16px; color: #334155; margin-bottom: 20px;">
                    Hello,
                </p>
                
                <p style="font-size: 16px; color: #334155; margin-bottom: 20px;">
                    You requested student pricing ($4.99) for AIE ResuMaker. Click the button below to verify your student status:
                </p>
                
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{verify_link}" 
                       style="background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%); 
                              color: white; padding: 14px 32px; text-decoration: none; 
                              border-radius: 8px; font-size: 16px; font-weight: bold; 
                              display: inline-block;">
                        ✅ Verify Student Status
                    </a>
                </div>
                
                <p style="font-size: 14px; color: #64748b; margin-bottom: 20px;">
                    Or copy and paste this link:
                </p>
                
                <p style="font-size: 13px; color: #475569; background: #f1f5f9; padding: 12px; border-radius: 6px; word-break: break-all;">
                    {verify_link}
                </p>
                
                <p style="font-size: 14px; color: #94a3b8; margin-top: 30px;">
                    This link expires in 24 hours. If you didn't request this, please ignore this email.
                </p>
                
                <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
                
                <p style="font-size: 12px; color: #94a3b8; text-align: center;">
                    AIE ResuMaker | Built in the US for US students and job seekers
                </p>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(html_body, 'html'))
        
        # Send email
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_FROM, to_email, msg.as_string())
        
        print(f"[EMAIL] Verification sent to {to_email}")
        return True
        
    except Exception as e:
        print(f"[EMAIL ERROR] Failed to send to {to_email}: {e}")
        return False

print(f"[AIE ResuMaker] Environment: {APP_ENV}")
print(f"[AIE ResuMaker] Stripe key prefix: {STRIPE_SECRET_KEY[:7]}..." if STRIPE_SECRET_KEY else "[AIE ResuMaker] WARNING: No Stripe secret key configured!")

from voice_api import router as voice_router, voice_sessions as voice_session_store

app = FastAPI(title="AIE ResuMaker", version="1.0")
app.include_router(voice_router)

@app.head("/")
@app.get("/healthz")
async def health_check():
    return {"status": "ok", "env": APP_ENV}

# O*NET API Configuration
ONET_API_KEY = os.getenv("ONET_API_KEY", "AYRvP-2AaMc-2BKqY-4UFda")
ONET_BASE_URL = "https://api-v2.onetcenter.org"

# Setup templates and static
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates_dir = os.path.join(BASE_DIR, "templates")
static_dir = os.path.join(BASE_DIR, "static")

# Mount static files
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Setup templates
templates = Jinja2Templates(directory=templates_dir)

# === Fix: Jinja2's _load_template uses (loader_ref, name) as cache key and
# passes globals (which may contain unhashable dicts/lists) into the cache.
# Monkey-patch _load_template to skip the cache lookup entirely.
import jinja2 as _jinja2
from jinja2.utils import LRUCache as _LRUCache

# Monkey-patch LRUCache.__getitem__ to avoid hash-based dict lookup
# which fails when cache keys contain unhashable types (dicts/lists).
_orig_lru_getitem = _LRUCache.__getitem__
def _safe_lru_getitem(self, key):
    with self._wlock:
        for k, v in self._mapping.items():
            if k == key:
                if self._queue[-1] != key:
                    try:
                        self._remove(key)
                    except ValueError:
                        pass
                    self._append(key)
                return v
        raise KeyError(key)
_LRUCache.__getitem__ = _safe_lru_getitem

# Also monkey-patch LRUCache.get so it uses our patched __getitem__
_LRUCache.get = lambda self, key, default=None: self[key] if key in [k for k in self._mapping.keys()] else default

# Monkey-patch _load_template to bypass cache lookup with unhashable keys
_orig_load_template = _jinja2.Environment._load_template
def _safe_load_template(self, name, globals):
    if self.loader is None:
        raise TypeError("no loader for this environment specified")
    # Always load fresh template, bypass cache.get()
    template = self.loader.load(self, name, self.make_globals(globals))
    if self.cache is not None:
        cache_key = (self.loader, name)  # use loader object directly instead of weakref
        self.cache[cache_key] = template
    return template
_jinja2.Environment._load_template = _safe_load_template



# Load Fortune 1000 companies
COMPANIES_FILE = os.path.join(BASE_DIR, "companies.json")
companies = []
try:
    with open(COMPANIES_FILE, 'r') as f:
        companies = json.load(f)
except:
    # Fallback list if file not found
    companies = [
        "Walmart", "Amazon", "Apple", "Costco", "JPMorgan Chase",
        "Microsoft", "Google", "Meta", "Berkshire Hathaway",
        "Johnson & Johnson", "UnitedHealth Group", "Exxon Mobil",
        "Chevron", "Home Depot", "Bank of America", "Wells Fargo",
        "Citigroup", "Goldman Sachs", "Morgan Stanley",
        "Procter & Gamble", "AT&T", "Verizon", "Comcast",
        "CVS Health", "Cardinal Health", "McKesson", "Cigna",
        "Anthem", "Lowe's", "Intel", "IBM", "Oracle", "Cisco",
        "Qualcomm", "HP", "Dell Technologies", "General Electric",
        "Ford Motor", "General Motors", "Tesla", "PepsiCo",
        "Coca-Cola", "Kraft Heinz", "Mondelez", "Colgate-Palmolive",
        "Nike", "Starbucks", "McDonald's", "Chipotle", "Target",
        "Best Buy", "Kroger", "Whole Foods", "Gap", "Nordstrom",
        "Macy's", "Kohl's", "Dollar General", "Walgreens",
        "Rite Aid", "HCA Healthcare", "Tenet Healthcare",
        "LabCorp", "Quest Diagnostics", "3M", "Honeywell",
        "Lockheed Martin", "Boeing", "Raytheon", "General Dynamics",
        "Northrop Grumman", "AECOM", "Jacobs Engineering", "Fluor",
        "Bechtel", "Turner Construction", "PulteGroup", "Lennar",
        "D.R. Horton", "Toll Brothers", "Prologis", "Simon Property",
        "Equinix", "Digital Realty", "American Tower", "Crown Castle"
    ]


# Load US cities data
CITIES_FILE = os.path.join(BASE_DIR, "cities.json")
us_cities_data = {"states": {}, "cities_by_state": {}}
try:
    with open(CITIES_FILE, 'r') as f:
        us_cities_data = json.load(f)
except:
    pass


# Load education data
UNIVERSITIES_FILE = os.path.join(BASE_DIR, "universities.json")
DEGREES_FILE = os.path.join(BASE_DIR, "degrees.json")
FIELDS_FILE = os.path.join(BASE_DIR, "fields_of_study.json")

universities = []
degrees = []
fields_of_study = []

try:
    with open(UNIVERSITIES_FILE, 'r') as f:
        universities = json.load(f)
except:
    pass

try:
    with open(DEGREES_FILE, 'r') as f:
        degrees = json.load(f)
except:
    pass

try:
    with open(FIELDS_FILE, 'r') as f:
        fields_of_study = json.load(f)
except:
    pass

# Load skill categories for auto-grouping
SKILL_CATEGORIES_FILE = os.path.join(BASE_DIR, "skill_categories.json")
skill_categories = {}
skill_to_category = {}

def _load_skill_categories():
    """Load skill categories from JSON and build reverse lookup map"""
    global skill_categories, skill_to_category
    try:
        with open(SKILL_CATEGORIES_FILE, 'r') as f:
            data = json.load(f)
            skill_categories = data.get("categories", {})
            # Build reverse lookup: skill_lower -> category_name
            for category, skills in skill_categories.items():
                if category == "Other":
                    continue
                for skill in skills:
                    skill_to_category[skill.lower()] = category
    except Exception as e:
        print(f"Warning: Could not load skill categories: {e}")
        skill_categories = {}
        skill_to_category = {}

_load_skill_categories()

def categorize_skills(skills_list):
    """Categorize a flat list of skills into grouped dict.
    
    Returns dict like {"Languages": ["Python", "JavaScript"], "Cloud & DevOps": ["AWS"]}
    Unmatched skills go into "Other" category.
    """
    if not skills_list:
        return {}
    
    grouped = {cat: [] for cat in skill_categories.keys() if cat != "Other"}
    grouped["Other"] = []
    matched = set()
    
    for user_skill in skills_list:
        user_skill_stripped = user_skill.strip()
        user_skill_lower = user_skill_stripped.lower()
        found = False
        
        # 1. Exact match in reverse lookup
        if user_skill_lower in skill_to_category:
            category = skill_to_category[user_skill_lower]
            if user_skill_lower not in matched:
                grouped[category].append(user_skill_stripped)
                matched.add(user_skill_lower)
            found = True
        
        # 2. Substring match: user_skill is contained in a known skill
        if not found:
            for known_skill, category in skill_to_category.items():
                if user_skill_lower in known_skill or known_skill in user_skill_lower:
                    if user_skill_lower not in matched:
                        grouped[category].append(user_skill_stripped)
                        matched.add(user_skill_lower)
                    found = True
                    break
        
        # 3. Word-by-word partial match for multi-word skills
        if not found and " " in user_skill_lower:
            words = user_skill_lower.split()
            for word in words:
                if len(word) > 2 and word in skill_to_category:
                    category = skill_to_category[word]
                    if user_skill_lower not in matched:
                        grouped[category].append(user_skill_stripped)
                        matched.add(user_skill_lower)
                    found = True
                    break
        
        if not found:
            grouped["Other"].append(user_skill_stripped)
    
    # Remove empty categories
    return {k: v for k, v in grouped.items() if v}

import secrets

# In-memory storage for demo (use Redis/DB in production)
resumes = {}

# Student verification tokens: {token: {email: str, verified: bool, expires: datetime}}
student_verifications = {}

# Referral tracking: {code: {created_by: str, created_at: datetime, visits: int, conversions: int}}
referral_codes = {}

# Referral visits log: [{code: str, ip: str, timestamp: datetime, converted: bool}]
referral_visits = []


def is_edu_email(email: str) -> bool:
    """Check if email is a .edu address"""
    return email.lower().endswith('.edu')


# Test mode bypass for student verification
TEST_EDU_EMAILS = ["studentprice@test.edu", "test@student.edu", "demo@edu.test"]


def is_test_edu_email(email: str) -> bool:
    """Check if email is a test .edu email for development"""
    return email.lower() in TEST_EDU_EMAILS


def generate_verification_token() -> str:
    """Generate a unique verification token"""
    return secrets.token_urlsafe(32)


@app.post("/api/verify-edu")
async def verify_edu_email(request: Request):
    """Send verification link to .edu email"""
    try:
        data = await request.json()
        email = data.get("email", "").strip().lower()

        if not email:
            return JSONResponse({"error": "Email is required"}, status_code=400)

        # Temporarily removed .edu check for testing
        # if not is_edu_email(email):
        #     return JSONResponse({"error": "Must be a .edu email address"}, status_code=400)

        # Test mode: auto-verify test emails
        if APP_ENV == "development" and is_test_edu_email(email):
            token = generate_verification_token()
            student_verifications[token] = {
                "email": email,
                "verified": True,
                "expires": datetime.now() + timedelta(hours=24)
            }
            return {
                "success": True,
                "message": "Test email auto-verified!",
                "dev_token": token,
                "auto_verified": True
            }

        # Generate verification token
        token = generate_verification_token()
        expires = datetime.now() + timedelta(hours=24)

        student_verifications[token] = {
            "email": email,
            "verified": False,
            "expires": expires
        }

        # Send verification email
        email_sent = send_verification_email(email, token, BASE_URL)

        if email_sent:
            return {
                "success": True,
                "message": "Verification link sent to your .edu email",
                "dev_token": token if APP_ENV == "development" else None
            }
        else:
            return {
                "success": True,
                "message": "Email configured. Check your inbox.",
                "dev_token": token if APP_ENV == "development" else None,
                "warning": "If email doesn't arrive, use dev token for testing"
            }

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/verify-student")
async def verify_student_token(request: Request, token: str = ""):
    """Verify student status via email link"""
    try:
        if not token:
            return HTMLResponse("""
            <html><body style="font-family: Arial; text-align: center; padding: 50px;">
                <h2 style="color: #dc2626;">❌ Missing Token</h2>
                <p>No verification token provided.</p>
            </body></html>
            """)

        if token not in student_verifications:
            return HTMLResponse("""
            <html><body style="font-family: Arial; text-align: center; padding: 50px;">
                <h2 style="color: #dc2626;">❌ Invalid Link</h2>
                <p>This verification link is invalid or has already been used.</p>
            </body></html>
            """)

        verification = student_verifications[token]

        # Check if expired
        if datetime.now() > verification["expires"]:
            del student_verifications[token]
            return HTMLResponse("""
            <html><body style="font-family: Arial; text-align: center; padding: 50px;">
                <h2 style="color: #dc2626;">⏰ Link Expired</h2>
                <p>This verification link has expired. Please request a new one.</p>
            </body></html>
            """)

        # Mark as verified
        verification["verified"] = True

        return HTMLResponse(f"""
        <html><body style="font-family: Arial; text-align: center; padding: 50px; background: #f0fdf4;">
            <h2 style="color: #16a34a;">✅ Student Status Verified!</h2>
            <p>Your .edu email <strong>{verification["email"]}</strong> has been verified.</p>
            <p>You can now close this tab and return to AIE ResuMaker to complete your student purchase.</p>
            <br>
            <a href="/build" style="background: #3b82f6; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px;">Return to AIE ResuMaker</a>
        </body></html>
        """)

    except Exception as e:
        return HTMLResponse(f"""
        <html><body style="font-family: Arial; text-align: center; padding: 50px;">
            <h2 style="color: #dc2626;">❌ Error</h2>
            <p>{str(e)}</p>
        </body></html>
        """)


@app.post("/api/verify-edu-code")
async def verify_edu_code(request: Request):
    """Legacy endpoint - now uses token verification"""
    try:
        data = await request.json()
        token = data.get("token", "").strip()

        if not token:
            return JSONResponse({"error": "Token is required"}, status_code=400)

        if token not in student_verifications:
            return JSONResponse({"error": "Invalid token"}, status_code=400)

        verification = student_verifications[token]

        if datetime.now() > verification["expires"]:
            del student_verifications[token]
            return JSONResponse({"error": "Token expired"}, status_code=400)

        if verification["verified"]:
            return {
                "success": True,
                "verified": True,
                "email": verification["email"]
            }

        return JSONResponse({"error": "Not verified yet. Check your email for the verification link."}, status_code=400)

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/verify-edu-status")
async def check_edu_status(email: str = "", token: str = ""):
    """Check if student email is verified"""
    try:
        # Check by token
        if token and token in student_verifications:
            verification = student_verifications[token]
            if verification["verified"] and datetime.now() <= verification["expires"]:
                return {
                    "verified": True,
                    "is_student": True,
                    "email": verification["email"]
                }

        # Check by email
        if email:
            for token_key, verification in student_verifications.items():
                if verification["email"] == email.lower() and verification["verified"]:
                    if datetime.now() <= verification["expires"]:
                        return {
                            "verified": True,
                            "is_student": True,
                            "email": email
                        }

        return {"verified": False, "is_student": is_edu_email(email)}

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/mockup", response_class=HTMLResponse)
async def mockup_page(request: Request):
    return templates.TemplateResponse(request=request, name="landing-mockup.html")

@app.get("/preview-mockup", response_class=HTMLResponse)
async def preview_mockup_page(request: Request):
    return templates.TemplateResponse(request=request, name="preview-mockup.html")

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse(request=request, name="landing.html")

@app.get("/voice_chat", response_class=HTMLResponse)
async def voice_chat_page(request: Request):
    """Direct voice chat page."""
    return templates.TemplateResponse(request=request, name="voice_chat.html")

@app.get("/build", response_class=HTMLResponse)
async def build_page(request: Request):
    ref_code = request.query_params.get("ref", "")
    mode = request.query_params.get("mode", "")
    voice_session = request.query_params.get("voice_session", "")
    
    # Track referral visit if ref code provided
    if ref_code and ref_code in referral_codes:
        referral_codes[ref_code]["visits"] += 1
        referral_visits.append({
            "code": ref_code,
            "ip": request.client.host if request.client else "unknown",
            "timestamp": datetime.now(),
            "converted": False
        })
        print(f"[REFERRAL] Visit tracked for code: {ref_code}, total visits: {referral_codes[ref_code]['visits']}")
    
    # Check for voice session data (new voice_sessions from voice_api)
    voice_data = None
    if voice_session and voice_session in voice_session_store:
        raw_data = voice_session_store[voice_session].get("data", {}).copy()
        # Parse JSON string fields into proper arrays
        for field in ["experience", "education"]:
            val = raw_data.get(field)
            if isinstance(val, str):
                try:
                    raw_data[field] = json.loads(val)
                except Exception:
                    raw_data[field] = []
        # Parse skills string into array
        if isinstance(raw_data.get("skills"), str):
            raw_data["skills"] = [s.strip() for s in raw_data["skills"].split(",") if s.strip()]
        voice_data = raw_data
        print(f"[BUILD] Loading voice session {voice_session}")
    
    # Determine mode
    if mode == "voice":
        return templates.TemplateResponse(request=request, name="voice_chat.html")
    elif mode == "form":
        return templates.TemplateResponse(request=request, name="index.html", context={
            "voice_data": json.dumps(voice_data) if voice_data else "null",
            "session_id": voice_session
        })
    
    # Auto-detect: mobile -> voice, desktop -> form
    user_agent = request.headers.get("user-agent", "").lower()
    is_mobile = any(device in user_agent for device in ["iphone", "android", "ipad", "mobile"])
    
    if is_mobile:
        return templates.TemplateResponse(request=request, name="voice_chat.html")
    else:
        ctx = {
            "voice_data": json.dumps(voice_data) if voice_data else "null",
            "session_id": voice_session
        }
        return templates.TemplateResponse(request=request, name="index.html", context=ctx)

@app.get("/debug/storage")
async def debug_storage():
    """Debug endpoint to check browser storage - returns a page that dumps localStorage"""
    html_content = """
    <!DOCTYPE html>
    <html>
    <head><title>Debug Storage</title></head>
    <body>
    <h1>Browser Storage Debug</h1>
    <pre id="output"></pre>
    <script>
        const data = localStorage.getItem('aie_resume_progress');
        const output = document.getElementById('output');
        if (data) {
            const parsed = JSON.parse(data);
            output.textContent = JSON.stringify(parsed, null, 2);
        } else {
            output.textContent = 'No saved data found in localStorage';
        }
    </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

@app.get("/terms", response_class=HTMLResponse)
async def terms(request: Request):
    """Terms of Service page"""
    return templates.TemplateResponse(request=request, name="terms.html")

@app.get("/api/industries")
async def get_industries():
    """Return list of industries for dropdown"""
    return {
        "industries": [
            "Technology", "Healthcare", "Finance", "Education",
            "Marketing", "Sales", "Operations", "Legal",
            "Engineering", "Design", "Human Resources",
            "Customer Service", "Manufacturing", "Retail"
        ]
    }

@app.get("/api/jobs/search")
async def search_jobs(query: str = ""):
    """Search job titles via O*NET API proxy"""
    if not query or len(query) < 2:
        return {"results": []}
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"{ONET_BASE_URL}/online/search",
                params={"keyword": query, "start": 1, "end": 20},
                headers={"X-API-Key": ONET_API_KEY},
                timeout=5.0
            )
            
            if response.status_code == 200:
                data = response.json()
                results = []
                if "occupation" in data:
                    for occ in data["occupation"]:
                        results.append({
                            "code": occ.get("code", ""),
                            "title": occ.get("title", ""),
                            "tags": occ.get("tags", [])
                        })
                return {"results": results[:10]}
            else:
                return {"results": [], "error": f"API returned {response.status_code}"}
                
    except Exception as e:
        return {"results": [], "error": str(e)}

@app.get("/api/jobs/{job_code}")
async def get_job_details(job_code: str):
    """Get detailed job information from O*NET"""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"{ONET_BASE_URL}/online/occupations/{job_code}",
                headers={"X-API-Key": ONET_API_KEY},
                timeout=5.0
            )
            
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"API returned {response.status_code}"}
                
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/jobs/{job_code}/summary")
async def get_job_summary(job_code: str):
    """Get job summary for resume suggestions"""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"{ONET_BASE_URL}/online/occupations/{job_code}/summary",
                headers={"X-API-Key": ONET_API_KEY},
                timeout=5.0
            )
            
            if response.status_code == 200:
                data = response.json()
                return {
                    "title": data.get("title", ""),
                    "description": data.get("description", ""),
                    "skills": data.get("skills", []),
                    "knowledge": data.get("knowledge", []),
                    "abilities": data.get("abilities", [])
                }
            else:
                return {"error": f"API returned {response.status_code}"}
                
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/build")
async def build_resume(
    full_name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(""),
    state: str = Form(""),
    city: str = Form(""),
    linkedin: str = Form(""),
    website: str = Form(""),
    summary: str = Form(""),
    experience: str = Form(""),
    education: str = Form(""),
    skills: str = Form(""),
    projects: str = Form(""),
    competencies: str = Form(""),
    community: str = Form(""),
    certifications: str = Form(""),
    industry: str = Form(""),
    template_style: str = Form("professional")
):
    """Build resume and return download URLs"""
    
    exp_list = json.loads(experience) if experience else []
    edu_list = json.loads(education) if education else []
    skills_list = [s.strip() for s in skills.split("|") if s.strip()]
    proj_list = json.loads(projects) if projects else []
    comp_list = json.loads(competencies) if competencies else []
    comm_list = json.loads(community) if community else []
    cert_list = json.loads(certifications) if certifications else []
    
    # Build location string from state and city
    location_parts = []
    if city: location_parts.append(city)
    if state: location_parts.append(state)
    location = ", ".join(location_parts) if location_parts else ""
    
    # Auto-categorize skills for grouped preview display
    skills_categorized = categorize_skills(skills_list)
    
    resume_data = {
        "full_name": full_name,
        "email": email,
        "phone": phone,
        "location": location,
        "linkedin": linkedin,
        "website": website,
        "summary": summary,
        "experience": exp_list,
        "education": edu_list,
        "skills": skills_list,
        "skills_categorized": skills_categorized,
        "projects": proj_list,
        "competencies": comp_list,
        "community": comm_list,
        "certifications": cert_list,
        "industry": industry,
        "template_style": template_style,
        "created_at": datetime.now().isoformat()
    }
    
    resume_id = f"{full_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    resumes[resume_id] = resume_data
    
    docx_path = generate_docx(resume_id, resume_data)
    pdf_path = generate_pdf(resume_id, resume_data)
    
    return {
        "success": True,
        "resume_id": resume_id,
        "download_url": f"/api/download/{resume_id}",
        "download_url_pdf": f"/api/download/{resume_id}?format=pdf",
        "preview_html": generate_preview_html(resume_data, template_style)
    }

def generate_docx(resume_id: str, data: dict):
    """Generate Word document"""
    doc = Document()
    
    sections = doc.sections[0]
    sections.top_margin = Inches(0.75)
    sections.bottom_margin = Inches(0.75)
    sections.left_margin = Inches(0.75)
    sections.right_margin = Inches(0.75)
    
    # Name - centered, uppercase, bold
    name = doc.add_paragraph()
    name_run = name.add_run(data["full_name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_parts = []
    if data.get("location"): contact_parts.append(data["location"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    
    if contact_parts:
        contact = doc.add_paragraph(" | ".join(contact_parts))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.runs[0].font.size = Pt(9)
        contact.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Horizontal line
    doc.add_paragraph("_" * 80).runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    
    # Professional Summary
    if data.get("summary"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL SUMMARY")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        summary = doc.add_paragraph(data["summary"])
        summary.paragraph_format.space_after = Pt(12)
        summary.runs[0].font.size = Pt(10)
    
    # Technical Skills
    if data.get("skills"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("TECHNICAL SKILLS")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        skills_text = ", ".join(data["skills"]) if isinstance(data["skills"], list) else str(data["skills"])
        skills = doc.add_paragraph(skills_text)
        skills.paragraph_format.space_after = Pt(12)
        skills.runs[0].font.size = Pt(10)
    
    # Professional Experience
    if data.get("experience"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL EXPERIENCE")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for job in data["experience"]:
            # Job title — Company | Location (dates right-aligned)
            job_header = doc.add_paragraph()
            job_header.paragraph_format.space_after = Pt(2)
            
            title_run = job_header.add_run(f"{job.get('title', '')}")
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            
            if job.get("company"):
                company_run = job_header.add_run(f" — {job['company']}")
                company_run.bold = True
                company_run.font.size = Pt(10.5)
            
            # Location and dates on next line
            location_parts = []
            if job.get("city"): location_parts.append(job["city"])
            if job.get("state"): location_parts.append(job["state"])
            location_str = ", ".join(location_parts)
            
            if location_str or job.get("dates") or job.get("phone") or job.get("address"):
                loc_date = doc.add_paragraph()
                loc_date.paragraph_format.space_after = Pt(2)
                
                if location_str:
                    loc_run = loc_date.add_run(location_str)
                    loc_run.italic = True
                    loc_run.font.size = Pt(9.5)
                
                if job.get("dates"):
                    if location_str:
                        spacer = loc_date.add_run("  |  ")
                    date_run = loc_date.add_run(job["dates"])
                    date_run.italic = True
                    date_run.font.size = Pt(9.5)
                    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
                if job.get("phone"):
                    if location_str or job.get("dates"):
                        spacer2 = loc_date.add_run("  |  ")
                    phone_run = loc_date.add_run(job["phone"])
                    phone_run.italic = True
                    phone_run.font.size = Pt(9.5)
                    phone_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            if job.get("address"):
                addr = doc.add_paragraph(job["address"])
                addr.paragraph_format.space_after = Pt(2)
                addr.runs[0].italic = True
                addr.runs[0].font.size = Pt(9.5)
                addr.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            # Description
            if job.get("description"):
                desc = doc.add_paragraph(job["description"])
                desc.paragraph_format.space_after = Pt(8)
                desc.paragraph_format.left_indent = Inches(0.2)
                desc.runs[0].font.size = Pt(10)
    
    # Projects
    if data.get("projects"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROJECTS")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for project in data["projects"]:
            if isinstance(project, dict):
                # Project name
                proj_name = doc.add_paragraph()
                proj_name_run = proj_name.add_run(project.get('name', ''))
                proj_name_run.bold = True
                proj_name_run.font.size = Pt(10.5)
                
                # Tech stack
                if project.get("tech"):
                    tech = doc.add_paragraph()
                    tech_run = tech.add_run(project["tech"])
                    tech_run.italic = True
                    tech_run.font.size = Pt(9)
                    tech_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
                # Description
                if project.get("description"):
                    desc = doc.add_paragraph(project["description"])
                    desc.paragraph_format.space_after = Pt(4)
                    desc.runs[0].font.size = Pt(10)
                
                # Result
                if project.get("result"):
                    result = doc.add_paragraph()
                    result_run = result.add_run(project["result"])
                    result_run.italic = True
                    result_run.font.size = Pt(9.5)
                    result_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    
    # Notable Competencies
    if data.get("competencies"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("NOTABLE COMPETENCIES")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for comp in data["competencies"]:
            if isinstance(comp, dict):
                label = comp.get('label', '') or comp.get('name', '')
                desc = comp.get('description', '')
                if label and desc:
                    comp_para = doc.add_paragraph()
                    label_run = comp_para.add_run(f"{label}: ")
                    label_run.bold = True
                    label_run.font.size = Pt(9.5)
                    desc_run = comp_para.add_run(desc)
                elif label:
                    comp_para = doc.add_paragraph()
                    label_run = comp_para.add_run(label)
                    label_run.bold = True
                    label_run.font.size = Pt(9.5)
                desc_run.font.size = Pt(9.5)
    
    # Education
    if data.get("education"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("EDUCATION")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for edu in data["education"]:
            # School — Location
            school = doc.add_paragraph()
            school_run = school.add_run(edu.get('school', ''))
            school_run.bold = True
            school_run.font.size = Pt(10.5)
            
            if edu.get("location"):
                loc_run = school.add_run(f" — {edu['location']}")
                loc_run.bold = True
                loc_run.font.size = Pt(10.5)
            
            # Degree, field, dates
            degree_parts = []
            if edu.get("degree"): degree_parts.append(edu["degree"])
            if edu.get("field"): degree_parts.append(edu["field"])
            if edu.get("dates"): degree_parts.append(edu["dates"])
            if degree_parts:
                degree = doc.add_paragraph(", ".join(degree_parts))
                degree.runs[0].italic = True
                degree.runs[0].font.size = Pt(9.5)
                degree.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            # Honors
            if edu.get("honors"):
                honors = doc.add_paragraph(edu["honors"])
                honors.runs[0].font.size = Pt(9)
                honors.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            school.paragraph_format.space_after = Pt(8)
    
    # Community Involvement
    if data.get("community"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("COMMUNITY INVOLVEMENT")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for item in data["community"]:
            if isinstance(item, dict):
                comm = doc.add_paragraph()
                event_run = comm.add_run(item.get('event', ''))
                event_run.font.size = Pt(10)
                if item.get("organization"):
                    org_run = comm.add_run(f" | {item['organization']}")
                    org_run.font.size = Pt(10)
    
    # Certifications
    if data.get("certifications"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("CERTIFICATIONS")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        
        for cert in data["certifications"]:
            if isinstance(cert, dict):
                cert_para = doc.add_paragraph()
                name_run = cert_para.add_run(cert.get('name', ''))
                name_run.font.size = Pt(10)
                if cert.get("organization"):
                    org_run = cert_para.add_run(f" — {cert['organization']}")
                    org_run.font.size = Pt(10)
                if cert.get("date"):
                    date_run = cert_para.add_run(f", {cert['date']}")
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # References
    refs = doc.add_paragraph("*References available upon request*")
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs.runs[0].italic = True
    refs.runs[0].font.size = Pt(9)
    refs.paragraph_format.space_before = Pt(16)
    
    output_dir = os.path.join(tempfile.gettempdir(), "resumes")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"{resume_id}.docx")
    doc.save(filepath)
    
    return filepath

def generate_preview_html(data: dict, template_style: str = "professional"):
    """Generate HTML preview of resume based on template style"""
    
    # Build contact line
    contact_parts = []
    if data.get("location"): contact_parts.append(data["location"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    
    contact_line = " | ".join(contact_parts) if contact_parts else ""
    
    # Template-specific CSS
    if template_style == "modern":
        css = _get_modern_css()
    elif template_style == "minimal":
        css = _get_minimal_css()
    else:  # professional (default)
        css = _get_professional_css()
    
    html = f"""
    <div class="resume-preview {template_style}">
        <style>{css}</style>
        <div class="header">
            <div class="name">{data['full_name']}</div>
            <div class="contact">{contact_line}</div>
        </div>
    """
    
    # Rest of the content generation remains the same
    # Professional Summary
    if data.get("summary"):
        html += f"""
        <div class="section-title">Professional Summary</div>
        <div class="summary">{data['summary']}</div>
        """
    
    # Technical Skills
    skills_data = data.get("skills_categorized") or data.get("skills", [])
    if skills_data:
        html += '<div class="section-title">Technical Skills</div>'
        if isinstance(skills_data, dict):
            html += '<div class="skills-grid">'
            for category, skill_list in skills_data.items():
                html += f'''
                <div>
                    <div class="skill-category">{category}</div>
                    <div class="skill-list">{", ".join(skill_list) if isinstance(skill_list, list) else skill_list}</div>
                </div>
                '''
            html += '</div>'
        else:
            html += '<div class="skills-grid">'
            html += f'<div class="skill-list">{", ".join(skills_data) if isinstance(skills_data, list) else skills_data}</div>'
            html += '</div>'
    
    # Professional Experience
    if data.get("experience"):
        html += '<div class="section-title">Professional Experience</div>'
        for job in data["experience"]:
            company = job.get('company', '')
            title = job.get('title', '')
            dates = job.get('dates', '')
            location = job.get('location', '')
            description = job.get('description', '')
            
            html += f'''
            <div style="margin-bottom: 10px;">
                <div class="job-header">
                    <span class="job-title-company">{title}{' — ' + company if company else ''}</span>
                    <span class="job-dates">{dates}</span>
                </div>
            '''
            
            # Build location from city/state if not already present
            if not location:
                location_parts = []
                if job.get('city'): location_parts.append(job['city'])
                if job.get('state'): location_parts.append(job['state'])
                location = ", ".join(location_parts)
            
            if location:
                html += f'<div class="job-location">{location}</div>'
            
            if job.get('phone'):
                html += f'<div class="job-contact">Phone: {job["phone"]}</div>'
            
            if job.get('address'):
                html += f'<div class="job-contact">{job["address"]}</div>'
            
            if description:
                if '\n' in description or '\r' in description:
                    lines = [line.strip() for line in description.replace('\r', '\n').split('\n') if line.strip()]
                    if lines:
                        html += '<div class="job-description"><ul>'
                        for line in lines:
                            if line.startswith('•') or line.startswith('-'):
                                line = line[1:].strip()
                            html += f'<li>{line}</li>'
                        html += '</ul></div>'
                else:
                    html += f'<div class="job-description">{description}</div>'
            
            html += '</div>'
    
    # Projects
    if data.get("projects"):
        html += '<div class="section-title">Projects</div>'
        projects = data["projects"] if isinstance(data["projects"], list) else []
        for project in projects:
            if isinstance(project, dict):
                name = project.get('name', '')
                tech = project.get('tech', '')
                description = project.get('description', '')
                result = project.get('result', '')
                
                html += f'''
                <div style="margin-bottom: 10px;">
                    <div class="project-name">{name}</div>
                '''
                if tech:
                    html += f'<div class="project-tech">{tech}</div>'
                if description:
                    html += f'<div>{description}</div>'
                if result:
                    html += f'<div class="project-result">{result}</div>'
                html += '</div>'
    
    # Notable Competencies
    if data.get("competencies"):
        html += '<div class="section-title">Notable Competencies</div>'
        competencies = data["competencies"] if isinstance(data["competencies"], list) else []
        html += '<div class="competencies-grid">'
        for comp in competencies:
            if isinstance(comp, dict):
                label = comp.get('label', '') or comp.get('name', '')
                desc = comp.get('description', '')
                if label and desc:
                    html += f'<div class="competency-item"><span class="competency-label">{label}:</span> {desc}</div>'
                elif label:
                    html += f'<div class="competency-item"><span class="competency-label">{label}</span></div>'
                else:
                    html += f'<div class="competency-item">{comp}</div>'
            else:
                html += f'<div class="competency-item">{comp}</div>'
        html += '</div>'
    
    # Education
    if data.get("education"):
        html += '<div class="section-title">Education</div>'
        for edu in data["education"]:
            school = edu.get('school', '')
            degree = edu.get('degree', '')
            field = edu.get('field', '')
            dates = edu.get('dates', '')
            location = edu.get('location', '')
            honors = edu.get('honors', '')
            
            html += '<div class="education-item">'
            html += f'<div class="school-name">{school}{' — ' + location if location else ''}</div>'
            
            degree_parts = []
            if degree: degree_parts.append(degree)
            if field: degree_parts.append(field)
            if dates: degree_parts.append(dates)
            if degree_parts:
                html += f'<div class="degree-info">{", ".join(degree_parts)}</div>'
            
            if honors:
                html += f'<div class="honors">{honors}</div>'
            
            html += '</div>'
    
    # Certifications
    if data.get("certifications"):
        html += '<div class="section-title">Certifications</div>'
        for cert in data["certifications"]:
            if isinstance(cert, dict):
                name = cert.get('name', '')
                organization = cert.get('organization', '')
                date = cert.get('date', '')
                
                cert_parts = [name]
                if organization: cert_parts.append(organization)
                if date: cert_parts.append(date)
                html += f'<div>{", ".join(cert_parts)}</div>'
    
    # Community
    if data.get("community"):
        html += '<div class="section-title">Community Involvement</div>'
        for comm in data["community"]:
            if isinstance(comm, dict):
                event = comm.get('event', '')
                organization = comm.get('organization', '')
                html += f'<div class="community-item"><strong>{event}</strong>{', ' + organization if organization else ''}</div>'
    
    html += "</div>"
    return html

def _get_professional_css():
    """Professional template CSS - traditional, clean"""
    return """
        .resume-preview.professional {
            font-family: 'Times New Roman', Times, serif;
            font-size: 10.5pt;
            line-height: 1.3;
            color: #000;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.5in;
        }
        .resume-preview.professional .header {
            text-align: center;
            margin-bottom: 12px;
        }
        .resume-preview.professional .name {
            font-size: 18pt;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-bottom: 4px;
        }
        .resume-preview.professional .contact {
            font-size: 9pt;
            color: #333;
        }
        .resume-preview.professional .section-title {
            font-size: 11pt;
            font-weight: bold;
            text-transform: uppercase;
            color: #2c5aa0;
            border-bottom: 1px solid #2c5aa0;
            margin-top: 14px;
            margin-bottom: 6px;
            padding-bottom: 2px;
        }
        .resume-preview.professional .summary {
            text-align: justify;
            margin-bottom: 8px;
        }
        .resume-preview.professional .skills-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 8px 20px;
            margin-bottom: 8px;
        }
        .resume-preview.professional .skill-category {
            font-weight: bold;
            font-size: 10pt;
        }
        .resume-preview.professional .skill-list {
            font-size: 9.5pt;
        }
        .resume-preview.professional .job-header {
            display: flex;
            justify-content: space-between;
            align-items: baseline;
        }
        .resume-preview.professional .job-title-company {
            font-weight: bold;
            font-size: 10.5pt;
        }
        .resume-preview.professional .job-dates {
            font-style: italic;
            font-size: 9.5pt;
            color: #333;
        }
        .resume-preview.professional .job-location {
            font-size: 9.5pt;
            font-style: italic;
            margin-bottom: 3px;
        }
        .resume-preview.professional .job-description {
            margin-left: 0;
            margin-bottom: 8px;
        }
        .resume-preview.professional .job-description ul {
            margin: 3px 0;
            padding-left: 20px;
        }
        .resume-preview.professional .job-description li {
            margin-bottom: 2px;
        }
    """

def _get_modern_css():
    """Modern template CSS - color accents, cleaner fonts"""
    return """
        .resume-preview.modern {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            font-size: 10pt;
            line-height: 1.4;
            color: #2d3748;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.5in;
        }
        .resume-preview.modern .header {
            border-left: 4px solid #4a90e2;
            padding-left: 15px;
            margin-bottom: 20px;
        }
        .resume-preview.modern .name {
            font-size: 22pt;
            font-weight: 700;
            color: #1a365d;
            margin-bottom: 6px;
            letter-spacing: -0.5px;
        }
        .resume-preview.modern .contact {
            font-size: 9.5pt;
            color: #4a5568;
        }
        .resume-preview.modern .section-title {
            font-size: 12pt;
            font-weight: 700;
            color: #4a90e2;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-top: 18px;
            margin-bottom: 8px;
            padding-bottom: 4px;
            border-bottom: 2px solid #e2e8f0;
        }
        .resume-preview.modern .summary {
            color: #4a5568;
            margin-bottom: 10px;
        }
        .resume-preview.modern .skills-grid {
            display: flex;
            flex-wrap: wrap;
            gap: 6px;
            margin-bottom: 10px;
        }
        .resume-preview.modern .skill-category {
            background: #ebf8ff;
            color: #2b6cb0;
            padding: 3px 10px;
            border-radius: 12px;
            font-size: 9pt;
            font-weight: 600;
        }
        .resume-preview.modern .skill-list {
            font-size: 9.5pt;
        }
        .resume-preview.modern .job-header {
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 4px;
        }
        .resume-preview.modern .job-title-company {
            font-weight: 700;
            font-size: 11pt;
            color: #1a365d;
        }
        .resume-preview.modern .job-dates {
            font-size: 9pt;
            color: #718096;
            background: #f7fafc;
            padding: 2px 8px;
            border-radius: 4px;
        }
        .resume-preview.modern .job-location {
            font-size: 9pt;
            color: #718096;
            margin-bottom: 4px;
        }
        .resume-preview.modern .job-description {
            margin-bottom: 12px;
        }
        .resume-preview.modern .job-description ul {
            margin: 4px 0;
            padding-left: 20px;
        }
        .resume-preview.modern .job-description li {
            margin-bottom: 3px;
            color: #4a5568;
        }
    """

def _get_minimal_css():
    """Minimal template CSS - barebones, whitespace-heavy"""
    return """
        .resume-preview.minimal {
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            font-size: 10pt;
            line-height: 1.5;
            color: #333;
            max-width: 7in;
            margin: 0 auto;
            padding: 0.75in;
        }
        .resume-preview.minimal .header {
            margin-bottom: 30px;
            padding-bottom: 15px;
            border-bottom: 1px solid #eee;
        }
        .resume-preview.minimal .name {
            font-size: 24pt;
            font-weight: 300;
            letter-spacing: 2px;
            margin-bottom: 8px;
            color: #000;
        }
        .resume-preview.minimal .contact {
            font-size: 9pt;
            color: #666;
            letter-spacing: 0.5px;
        }
        .resume-preview.minimal .section-title {
            font-size: 10pt;
            font-weight: 400;
            text-transform: uppercase;
            letter-spacing: 2px;
            color: #999;
            margin-top: 25px;
            margin-bottom: 10px;
        }
        .resume-preview.minimal .summary {
            color: #555;
            margin-bottom: 10px;
        }
        .resume-preview.minimal .skills-grid {
            margin-bottom: 10px;
        }
        .resume-preview.minimal .skill-category {
            font-weight: 400;
            font-size: 9.5pt;
            color: #666;
        }
        .resume-preview.minimal .skill-list {
            font-size: 9.5pt;
            color: #666;
        }
        .resume-preview.minimal .job-header {
            margin-bottom: 4px;
        }
        .resume-preview.minimal .job-title-company {
            font-weight: 500;
            font-size: 11pt;
        }
        .resume-preview.minimal .job-dates {
            font-size: 9pt;
            color: #999;
            float: right;
        }
        .resume-preview.minimal .job-location {
            font-size: 9pt;
            color: #999;
            margin-bottom: 6px;
        }
        .resume-preview.minimal .job-description {
            margin-bottom: 15px;
            color: #555;
        }
        .resume-preview.minimal .job-description ul {
            margin: 5px 0;
            padding-left: 20px;
        }
        .resume-preview.minimal .job-description li {
            margin-bottom: 4px;
        }
    """
    
    # Professional Summary
    if data.get("summary"):
        html += f"""
        <div class="section-title">Professional Summary</div>
        <div class="summary">{data['summary']}</div>
        """
    
    # Technical Skills
    skills_data = data.get("skills_categorized") or data.get("skills", [])
    if skills_data:
        html += '<div class="section-title">Technical Skills</div>'
        # Group skills by category if provided, otherwise show as list
        if isinstance(skills_data, dict):
            html += '<div class="skills-grid">'
            for category, skill_list in skills_data.items():
                html += f'''
                <div>
                    <div class="skill-category">{category}</div>
                    <div class="skill-list">{', '.join(skill_list) if isinstance(skill_list, list) else skill_list}</div>
                </div>
                '''
            html += '</div>'
        else:
            # Simple list of skills
            html += '<div class="skills-grid">'
            html += f'<div class="skill-list">{', '.join(skills_data) if isinstance(skills_data, list) else skills_data}</div>'
            html += '</div>'
    
    # Professional Experience
    if data.get("experience"):
        html += '<div class="section-title">Professional Experience</div>'
        for job in data["experience"]:
            company = job.get('company', '')
            title = job.get('title', '')
            dates = job.get('dates', '')
            location = job.get('location', '')
            description = job.get('description', '')
            
            # Format: Title — Company | Location <span style="float:right">Dates</span>
            html += f'''
            <div style="margin-bottom: 10px;">
                <div class="job-header">
                    <span class="job-title-company">{title}{' — ' + company if company else ''}</span>
                    <span class="job-dates">{dates}</span>
                </div>
            '''
            
            # Build location from city/state if not already present
            if not location:
                location_parts = []
                if job.get('city'): location_parts.append(job['city'])
                if job.get('state'): location_parts.append(job['state'])
                location = ", ".join(location_parts)
            
            if location:
                html += f'<div class="job-location">{location}</div>'
            
            if job.get('phone'):
                html += f'<div class="job-contact">Phone: {job["phone"]}</div>'
            
            if job.get('address'):
                html += f'<div class="job-contact">{job["address"]}</div>'
            
            if description:
                # Convert description to bullet points if it contains line breaks
                if '\n' in description or '\r' in description:
                    lines = [line.strip() for line in description.replace('\r', '\n').split('\n') if line.strip()]
                    if lines:
                        html += '<div class="job-description"><ul>'
                        for line in lines:
                            if line.startswith('•') or line.startswith('-'):
                                line = line[1:].strip()
                            html += f'<li>{line}</li>'
                        html += '</ul></div>'
                else:
                    html += f'<div class="job-description">{description}</div>'
            
            html += '</div>'
    
    # Projects
    if data.get("projects"):
        html += '<div class="section-title">Projects</div>'
        projects = data["projects"] if isinstance(data["projects"], list) else []
        for project in projects:
            if isinstance(project, dict):
                name = project.get('name', '')
                tech = project.get('tech', '')
                description = project.get('description', '')
                result = project.get('result', '')
                
                html += f'''
                <div style="margin-bottom: 10px;">
                    <div class="project-name">{name}</div>
                '''
                if tech:
                    html += f'<div class="project-tech">{tech}</div>'
                if description:
                    html += f'<div>{description}</div>'
                if result:
                    html += f'<div class="project-result">{result}</div>'
                html += '</div>'
    
    # Notable Competencies
    if data.get("competencies"):
        html += '<div class="section-title">Notable Competencies</div>'
        competencies = data["competencies"] if isinstance(data["competencies"], list) else []
        html += '<div class="competencies-grid">'
        for comp in competencies:
            if isinstance(comp, dict):
                label = comp.get('label', '') or comp.get('name', '')
                desc = comp.get('description', '')
                if label and desc:
                    html += f'<div class="competency-item"><span class="competency-label">{label}:</span> {desc}</div>'
                elif label:
                    html += f'<div class="competency-item"><span class="competency-label">{label}</span></div>'
                else:
                    html += f'<div class="competency-item">{comp}</div>'
            else:
                html += f'<div class="competency-item">{comp}</div>'
        html += '</div>'
    
    # Education
    if data.get("education"):
        html += '<div class="section-title">Education</div>'
        for edu in data["education"]:
            school = edu.get('school', '')
            degree = edu.get('degree', '')
            field = edu.get('field', '')
            dates = edu.get('dates', '')
            location = edu.get('location', '')
            honors = edu.get('honors', '')
            
            html += '<div class="education-item">'
            # Format: School — Location
            html += f'<div class="school-name">{school}{' — ' + location if location else ''}</div>'
            
            # Degree info
            degree_parts = []
            if degree: degree_parts.append(degree)
            if field: degree_parts.append(field)
            if dates: degree_parts.append(dates)
            if degree_parts:
                html += f'<div class="degree-info">{', '.join(degree_parts)}</div>'
            
            if honors:
                html += f'<div class="honors">{honors}</div>'
            
            html += '</div>'
    
    # Community Involvement
    if data.get("community"):
        html += '<div class="section-title">Community Involvement</div>'
        community = data["community"] if isinstance(data["community"], list) else []
        for item in community:
            if isinstance(item, dict):
                event = item.get('event', '')
                org = item.get('organization', '')
                html += f'<div class="community-item">{event}{' | ' + org if org else ''}</div>'
            else:
                html += f'<div class="community-item">{item}</div>'
    
    # Certifications
    if data.get("certifications"):
        html += '<div class="section-title">Certifications</div>'
        certs = data["certifications"] if isinstance(data["certifications"], list) else []
        for cert in certs:
            if isinstance(cert, dict):
                name = cert.get('name', '')
                org = cert.get('organization', '')
                date = cert.get('date', '')
                html += f'<div>{name}{' — ' + org if org else ''}{', ' + date if date else ''}</div>'
            else:
                html += f'<div>{cert}</div>'
    
    html += '<div class="references">*References available upon request*</div>'
    html += "</div>"
    return html

@app.get("/api/download/{resume_id}")
async def download_resume(resume_id: str):
    """Download resume as .docx"""
    filepath = os.path.join(tempfile.gettempdir(), "resumes", f"{resume_id}.docx")
    if os.path.exists(filepath):
        return FileResponse(filepath, filename=f"resume_{resume_id}.docx")
    return JSONResponse({"error": "Resume not found"}, status_code=404)

@app.get("/api/companies/search")
async def search_companies(query: str = ""):
    """Search Fortune 1000 companies"""
    if not query or len(query) < 2:
        return {"results": []}
    
    query_lower = query.lower()
    results = []
    for company in companies:
        if query_lower in company.lower():
            results.append(company)
    
    return {"results": results[:10]}  # Limit to 10 suggestions

@app.get("/api/states")
async def get_states():
    """Return list of US states"""
    return {"states": us_cities_data.get("states", {})}

@app.get("/api/cities")
async def get_cities(state: str = ""):
    """Return cities for a given state"""
    if not state:
        return {"cities": []}
    
    cities = us_cities_data.get("cities_by_state", {}).get(state, [])
    return {"cities": cities}


@app.get("/api/universities")
async def get_universities(q: str = ""):
    """Search universities by name"""
    if not q:
        return {"results": universities[:50]}
    
    query_lower = q.lower()
    results = [u for u in universities if query_lower in u.lower()]
    return {"results": results[:10]}

@app.get("/api/degrees")
async def get_degrees():
    """Return list of degree types"""
    return {"degrees": degrees}

@app.get("/api/fields")
async def get_fields(q: str = ""):
    """Search fields of study"""
    if not q:
        return {"results": fields_of_study[:50]}
    
    query_lower = q.lower()
    results = [f for f in fields_of_study if query_lower in f.lower()]
    return {"results": results[:10]}


# Stripe Payment Endpoints

@app.post("/api/referral/create")
async def create_referral(request: Request):
    """Create a new referral code for a user"""
    try:
        data = await request.json()
        user_email = data.get("email", "anonymous")
        
        # Generate unique referral code
        code = secrets.token_urlsafe(8).upper()[:10]
        
        # Store referral code
        referral_codes[code] = {
            "created_by": user_email,
            "created_at": datetime.now(),
            "visits": 0,
            "conversions": 0,
            "reward_unlocked": False
        }
        
        return {
            "success": True,
            "code": code,
            "link": f"{BASE_URL}/build?ref={code}"
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/referral/track")
async def track_referral_visit(request: Request):
    """Track when a referral link is visited"""
    try:
        data = await request.json()
        code = data.get("code", "")
        ip_address = request.client.host if request.client else "unknown"
        
        if not code or code not in referral_codes:
            return JSONResponse({"error": "Invalid referral code"}, status_code=400)
        
        # Record visit
        referral_codes[code]["visits"] += 1
        referral_visits.append({
            "code": code,
            "ip": ip_address,
            "timestamp": datetime.now(),
            "converted": False
        })
        
        return {
            "success": True,
            "visits": referral_codes[code]["visits"],
            "conversions": referral_codes[code]["conversions"],
            "reward_unlocked": referral_codes[code]["reward_unlocked"]
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/referral/stats/{code}")
async def get_referral_stats(code: str):
    """Get referral statistics for a code"""
    if code not in referral_codes:
        return JSONResponse({"error": "Referral code not found"}, status_code=404)
    
    stats = referral_codes[code]
    return {
        "code": code,
        "visits": stats["visits"],
        "conversions": stats["conversions"],
        "reward_unlocked": stats["reward_unlocked"],
        "link": f"{BASE_URL}/build?ref={code}"
    }

@app.post("/api/referral/convert")
async def convert_referral(request: Request):
    """Mark a referral as converted (friend made a purchase)"""
    try:
        data = await request.json()
        code = data.get("code", "")
        
        if not code or code not in referral_codes:
            return JSONResponse({"error": "Invalid referral code"}, status_code=400)
        
        # Mark conversion
        referral_codes[code]["conversions"] += 1
        referral_codes[code]["reward_unlocked"] = True
        
        # Update visit log
        for visit in reversed(referral_visits):
            if visit["code"] == code and not visit["converted"]:
                visit["converted"] = True
                break
        
        return {
            "success": True,
            "conversions": referral_codes[code]["conversions"],
            "reward_unlocked": True
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/stripe/config")
async def get_stripe_config():
    """Return Stripe publishable key to frontend"""
    return {
        "publishableKey": STRIPE_PUBLISHABLE_KEY,
        "priceId": STRIPE_PRICE_ID,
        "discountPriceId": STRIPE_STUDENT_PRICE_ID,
        "studentPriceId": STRIPE_STUDENT_PRICE_ID
    }

@app.post("/api/create-checkout-session")
async def create_checkout_session(request: Request):
    """Create Stripe Checkout session for premium resume download"""
    try:
        data = await request.json()
        resume_id = data.get("resume_id", "")
        tier = data.get("tier", "regular")  # regular, discount, student
        student_token = data.get("student_token", "")
        
        if not resume_id:
            return JSONResponse({"error": "Resume ID required"}, status_code=400)
        
        # Verify student status if requesting student price
        if tier == "student":
            verified = False
            
            # Check token
            if student_token and student_token in student_verifications:
                verification = student_verifications[student_token]
                if verification["verified"] and datetime.now() <= verification["expires"]:
                    verified = True
            
            if not verified:
                return JSONResponse({
                    "error": "Student verification required. Please verify your .edu email first."
                }, status_code=403)
        
        # Map tier to price ID
        if tier == "student" or tier == "discount":
            price_id = STRIPE_STUDENT_PRICE_ID
        else:
            price_id = STRIPE_PRICE_ID
        
        # Create checkout session
        # Check if user came from referral
        ref_code = data.get("referral_code", "")
        
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price": price_id,
                "quantity": 1,
            }],
            mode="payment",
            success_url=f"{BASE_URL}/success?session_id={{CHECKOUT_SESSION_ID}}&resume_id={resume_id}",
            cancel_url=f"{BASE_URL}/cancel?resume_id={resume_id}",
            metadata={
                "resume_id": resume_id,
                "tier": tier,
                "referral_code": ref_code
            }
        )
        
        return {"sessionId": session.id, "url": session.url}
    
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/success")
async def payment_success(request: Request, session_id: str = "", resume_id: str = "", ref: str = ""):
    """Handle successful payment - show download page"""
    try:
        # Verify payment with Stripe
        session = stripe.checkout.Session.retrieve(session_id)
        
        if session.payment_status == "paid":
            # Payment confirmed - mark resume as premium
            if resume_id in resumes:
                resumes[resume_id]["paid"] = True
            
            # Track referral conversion
            ref_code = ref or session.get("metadata", {}).get("referral_code", "")
            if ref_code and ref_code in referral_codes:
                referral_codes[ref_code]["conversions"] += 1
                referral_codes[ref_code]["reward_unlocked"] = True
                print(f"[REFERRAL] Conversion tracked for code: {ref_code}")
            
            return templates.TemplateResponse(
                request=request,
                name="success.html",
                context={
                    "resume_id": resume_id,
                    "download_url": f"/api/download/{resume_id}"
                }
            )
        else:
            return HTMLResponse("Payment not completed. Please try again.")
    
    except Exception as e:
        return HTMLResponse(f"Error verifying payment: {str(e)}")

@app.get("/cancel")
async def payment_cancel(request: Request, resume_id: str = ""):
    """Handle cancelled payment"""
    return templates.TemplateResponse(
        request=request,
        name="cancel.html",
        context={"resume_id": resume_id}
    )

@app.post("/api/webhook")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events"""
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature", "")
    webhook_secret = STRIPE_WEBHOOK_SECRET
    
    try:
        if webhook_secret:
            event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        else:
            # For testing without webhook secret
            event = json.loads(payload)
        
        # Handle successful payment
        if event.get("type") == "checkout.session.completed":
            session = event["data"]["object"]
            resume_id = session.get("metadata", {}).get("resume_id", "")
            ref_code = session.get("metadata", {}).get("referral_code", "")
            
            if resume_id and resume_id in resumes:
                resumes[resume_id]["paid"] = True
                print(f"Payment confirmed for resume: {resume_id}")
            
            # Track referral conversion from webhook
            if ref_code and ref_code in referral_codes:
                referral_codes[ref_code]["conversions"] += 1
                referral_codes[ref_code]["reward_unlocked"] = True
                print(f"[REFERRAL] Webhook conversion tracked for code: {ref_code}")
        
        return {"status": "success"}
    
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

@app.get("/api/check-payment/{resume_id}")
async def check_payment(resume_id: str):
    """Check if resume has been paid for"""
    resume = resumes.get(resume_id, {})
    return {"paid": resume.get("paid", False)}


# Anti-Piracy: Preview Rendering with Watermark
from playwright.async_api import async_playwright
from PIL import Image, ImageDraw, ImageFont

async def render_preview_to_image(html_content: str, resume_id: str) -> bytes:
    """Render preview HTML to watermarked PNG image"""
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page(viewport={"width": 816, "height": 1056})
            
            # Set content
            await page.set_content(html_content)
            await page.wait_for_load_state("networkidle")
            
            # Screenshot
            screenshot = await page.screenshot(type="png")
            await browser.close()
            
            # Open with PIL
            img = Image.open(BytesIO(screenshot))
            
            # Add watermark
            draw = ImageDraw.Draw(img)
            
            # Try to get a font, fallback to default
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
                small_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
            except:
                font = ImageFont.load_default()
                small_font = ImageFont.load_default()
            
            # Watermark text
            watermark_text = "AIE ResuMaker SAMPLE"
            sub_text = "To get the highest quality format, click Purchase"
            
            # Calculate position for diagonal watermark
            bbox = draw.textbbox((0, 0), watermark_text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # Draw multiple diagonal watermarks
            img_width, img_height = img.size
            for i in range(-2, 4):
                x = (i * 300) % (img_width + text_width) - text_width
                y = int(img_height * 0.3 + (i * 150))
                
                # Semi-transparent white background
                overlay = Image.new('RGBA', img.size, (255, 255, 255, 0))
                overlay_draw = ImageDraw.Draw(overlay)
                overlay_draw.text((x, y), watermark_text, font=font, fill=(200, 0, 0, 80))
                
                # Composite
                img = Image.alpha_composite(img.convert('RGBA'), overlay).convert('RGB')
                draw = ImageDraw.Draw(img)
            
            # Add bottom message
            bottom_box_height = 40
            overlay = Image.new('RGBA', img.size, (255, 255, 255, 0))
            overlay_draw = ImageDraw.Draw(overlay)
            overlay_draw.rectangle([0, img_height - bottom_box_height, img_width, img_height], 
                                   fill=(0, 0, 0, 180))
            img = Image.alpha_composite(img.convert('RGBA'), overlay).convert('RGB')
            draw = ImageDraw.Draw(img)
            
            bbox = draw.textbbox((0, 0), sub_text, font=small_font)
            sub_width = bbox[2] - bbox[0]
            sub_x = (img_width - sub_width) // 2
            draw.text((sub_x, img_height - 30), sub_text, font=small_font, fill=(255, 255, 255))
            
            # Compress to lower quality
            output = BytesIO()
            img.save(output, format='PNG', quality=50, optimize=True)
            output.seek(0)
            
            return output.getvalue()
            
    except Exception as e:
        print(f"Error rendering preview image: {e}")
        return None

@app.post("/api/preview-timer")
async def get_timed_preview(request: Request):
    """Get preview with 3-second clean timer"""
    try:
        data = await request.json()
        resume_id = data.get("resume_id", "")
        
        if not resume_id or resume_id not in resumes:
            return JSONResponse({"error": "Resume not found"}, status_code=404)
        
        resume_data = resumes[resume_id]
        template_style = resume_data.get("template_style", "professional")
        
        # Generate watermarked image with correct theme
        preview_html = generate_preview_html(resume_data, template_style)
        image_bytes = await render_preview_to_image(preview_html, resume_id)
        
        if image_bytes:
            # Convert to base64 for frontend
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            return {
                "success": True,
                "clean_html": preview_html,
                "watermarked_image": f"data:image/png;base64,{image_base64}"
            }
        else:
            return JSONResponse({"error": "Failed to generate preview"}, status_code=500)
            
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

def generate_pdf(resume_id: str, data: dict):
    """Generate PDF document"""
    output_dir = os.path.join(tempfile.gettempdir(), "resumes")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"{resume_id}.pdf")
    
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#000000'),
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Times-Bold'
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#333333'),
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=colors.HexColor('#2c5aa0'),
        spaceAfter=6,
        spaceBefore=14,
        fontName='Times-Bold',
        borderColor=colors.HexColor('#2c5aa0'),
        borderWidth=1,
        borderPadding=2,
        leftIndent=0,
        rightIndent=0
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        fontName='Times-Roman'
    )
    
    # Name
    story.append(Paragraph(data['full_name'].upper(), name_style))
    
    # Contact info
    contact_parts = []
    if data.get("location"): contact_parts.append(data["location"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))
    
    # Horizontal line
    story.append(Spacer(1, 12))
    
    # Professional Summary
    if data.get("summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(Paragraph(data["summary"], normal_style))
    
    # Technical Skills
    if data.get("skills"):
        story.append(Paragraph("TECHNICAL SKILLS", section_style))
        skills_text = ", ".join(data["skills"]) if isinstance(data["skills"], list) else str(data["skills"])
        story.append(Paragraph(skills_text, normal_style))
    
    # Professional Experience
    if data.get("experience"):
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_style))
        for job in data["experience"]:
            # Job title and company
            title_text = f"<b>{job.get('title', '')}</b>"
            if job.get("company"):
                title_text += f" — <b>{job['company']}</b>"
            story.append(Paragraph(title_text, ParagraphStyle('JobTitle', parent=normal_style, fontSize=10.5, fontName='Times-Bold', spaceAfter=2)))
            
            # Location and dates
            location_parts = []
            if job.get("city"): location_parts.append(job["city"])
            if job.get("state"): location_parts.append(job["state"])
            location_str = ", ".join(location_parts)
            
            loc_date_parts = []
            if location_str: loc_date_parts.append(f"<i>{location_str}</i>")
            if job.get("dates"): loc_date_parts.append(f"<i>{job['dates']}</i>")
            if job.get("phone"): loc_date_parts.append(f"<i>{job['phone']}</i>")
            
            if loc_date_parts:
                story.append(Paragraph("  |  ".join(loc_date_parts), ParagraphStyle('JobMeta', parent=normal_style, fontSize=9.5, textColor=colors.HexColor('#666666'))))
            
            # Description
            if job.get("description"):
                story.append(Paragraph(job["description"], ParagraphStyle('JobDesc', parent=normal_style, leftIndent=14)))
    
    # Projects
    if data.get("projects"):
        story.append(Paragraph("PROJECTS", section_style))
        for project in data["projects"]:
            if isinstance(project, dict):
                story.append(Paragraph(f"<b>{project.get('name', '')}</b>", ParagraphStyle('ProjectTitle', parent=normal_style, fontSize=10.5, fontName='Times-Bold')))
                if project.get("tech"):
                    story.append(Paragraph(f"<i>{project['tech']}</i>", ParagraphStyle('ProjectTech', parent=normal_style, fontSize=9, textColor=colors.HexColor('#666666'))))
                if project.get("description"):
                    story.append(Paragraph(project["description"], normal_style))
                if project.get("result"):
                    story.append(Paragraph(f"<i>{project['result']}</i>", ParagraphStyle('ProjectResult', parent=normal_style, fontSize=9.5, textColor=colors.HexColor('#2c5aa0'))))
    
    # Notable Competencies
    if data.get("competencies"):
        story.append(Paragraph("NOTABLE COMPETENCIES", section_style))
        for comp in data["competencies"]:
            if isinstance(comp, dict):
                label = comp.get('label', '') or comp.get('name', '')
                desc = comp.get('description', '')
                if label and desc:
                    story.append(Paragraph(f"<b>{label}:</b> {desc}", normal_style))
                elif label:
                    story.append(Paragraph(f"<b>{label}</b>", normal_style))
    
    # Education
    if data.get("education"):
        story.append(Paragraph("EDUCATION", section_style))
        for edu in data["education"]:
            school_text = f"<b>{edu.get('school', '')}</b>"
            if edu.get("location"):
                school_text += f" — <b>{edu['location']}</b>"
            story.append(Paragraph(school_text, ParagraphStyle('School', parent=normal_style, fontSize=10.5, fontName='Times-Bold')))
            
            degree_parts = []
            if edu.get("degree"): degree_parts.append(edu["degree"])
            if edu.get("field"): degree_parts.append(edu["field"])
            if edu.get("dates"): degree_parts.append(edu["dates"])
            if degree_parts:
                story.append(Paragraph(f"<i>{', '.join(degree_parts)}</i>", ParagraphStyle('Degree', parent=normal_style, fontSize=9.5, textColor=colors.HexColor('#666666'))))
            
            if edu.get("honors"):
                story.append(Paragraph(edu["honors"], ParagraphStyle('Honors', parent=normal_style, fontSize=9, textColor=colors.HexColor('#666666'))))
    
    # Community Involvement
    if data.get("community"):
        story.append(Paragraph("COMMUNITY INVOLVEMENT", section_style))
        for item in data["community"]:
            if isinstance(item, dict):
                text = item.get("event", "")
                if item.get("organization"):
                    text += f" | {item['organization']}"
                story.append(Paragraph(text, normal_style))
    
    # Certifications
    if data.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", section_style))
        for cert in data["certifications"]:
            if isinstance(cert, dict):
                text = cert.get("name", "")
                if cert.get("organization"):
                    text += f" — {cert['organization']}"
                if cert.get("date"):
                    text += f", {cert['date']}"
                story.append(Paragraph(text, normal_style))
    
    # References
    story.append(Spacer(1, 16))
    story.append(Paragraph("*References available upon request*", ParagraphStyle('References', parent=normal_style, alignment=TA_CENTER, fontSize=9, fontName='Times-Italic')))
    
    doc.build(story)
    return filepath


# Session storage for resume states (UUID -> ResumeState)
resume_sessions: Dict[str, "ResumeState"] = {}


class ResumeState:
    """Tracks the user's resume building progress with confidence scores."""
    
    def __init__(self, session_id: str = ""):
        self.session_id = session_id or secrets.token_urlsafe(16)
        self.created_at = datetime.now().isoformat()
        self.last_updated = self.created_at
        
        # Core fields with confidence scores (0.0 - 1.0)
        self.full_name = ""
        self.full_name_confidence = 0.0
        
        self.email = ""
        self.email_confidence = 0.0
        
        self.phone = ""
        self.phone_confidence = 0.0
        
        self.job_title = ""
        self.job_title_confidence = 0.0
        
        self.experience_level = "entry"  # entry|mid|senior|executive
        self.experience_level_confidence = 0.0
        
        self.summary = ""
        self.summary_confidence = 0.0
        
        # Arrays
        self.experience: List[Dict] = []
        self.education: List[Dict] = []
        self.skills: List[str] = []
        self.projects: List[Dict] = []
        self.competencies: List[Dict] = []
        self.community: List[Dict] = []
        self.certifications: List[Dict] = []
        
        # Tracking
        self.turn_count = 0  # How many voice interactions
        self.fields_filled = set()  # Which fields have data
        self.fields_missing = {"full_name", "email", "phone", "job_title", 
                               "experience", "education", "skills", "summary"}
    
    def update_field(self, field: str, value: any, confidence: float = 0.8):
        """Update a field and track its confidence."""
        if field in ["experience", "education", "projects", "competencies", 
                      "community", "certifications", "skills"]:
            # Array fields
            if isinstance(value, list) and len(value) > 0:
                setattr(self, field, value)
                self.fields_filled.add(field)
                self.fields_missing.discard(field)
        elif value and str(value).strip():
            # Scalar fields
            setattr(self, field, value)
            setattr(self, f"{field}_confidence", confidence)
            self.fields_filled.add(field)
            self.fields_missing.discard(field)
        
        self.last_updated = datetime.now().isoformat()
    
    def add_experience(self, job: Dict):
        """Add or update a job entry."""
        # Check if this is an update to existing job
        for existing in self.experience:
            if existing.get("company") == job.get("company") and \
               existing.get("title") == job.get("title"):
                # Merge missing fields
                for key in ["dates", "description"]:
                    if not existing.get(key) and job.get(key):
                        existing[key] = job[key]
                return
        
        # New job
        self.experience.append(job)
        self.fields_filled.add("experience")
        self.fields_missing.discard("experience")
        self.last_updated = datetime.now().isoformat()
    
    def add_education(self, edu: Dict):
        """Add or update education entry."""
        for existing in self.education:
            if existing.get("school") == edu.get("school"):
                for key in ["degree", "field", "dates"]:
                    if not existing.get(key) and edu.get(key):
                        existing[key] = edu[key]
                return
        
        self.education.append(edu)
        self.fields_filled.add("education")
        self.fields_missing.discard("education")
        self.last_updated = datetime.now().isoformat()
    
    def add_skills(self, new_skills: List[str]):
        """Add skills, deduplicate."""
        existing = set(self.skills)
        for skill in new_skills:
            skill_clean = skill.strip()
            if skill_clean and skill_clean.lower() not in [s.lower() for s in existing]:
                self.skills.append(skill_clean)
                existing.add(skill_clean.lower())
        
        if self.skills:
            self.fields_filled.add("skills")
            self.fields_missing.discard("skills")
        self.last_updated = datetime.now().isoformat()
    
    def infer_experience_level(self):
        """Auto-calculate experience level from work history."""
        total_years = 0
        for job in self.experience:
            dates = job.get("dates", "")
            # Try to extract years from "2020 - 2023" or "2020-2023"
            import re
            years = re.findall(r'\d{4}', dates)
            if len(years) >= 2:
                total_years += int(years[1]) - int(years[0])
            elif len(years) == 1 and "present" in dates.lower():
                total_years += 2026 - int(years[0])  # Current year
        
        if total_years >= 10:
            self.experience_level = "executive"
        elif total_years >= 5:
            self.experience_level = "senior"
        elif total_years >= 2:
            self.experience_level = "mid"
        else:
            self.experience_level = "entry"
        
        self.experience_level_confidence = min(0.9, 0.5 + total_years * 0.05)
        self.fields_filled.add("experience_level")
        self.fields_missing.discard("experience_level")
    
    def to_dict(self) -> Dict:
        """Convert to dict for JSON serialization."""
        return {
            "session_id": self.session_id,
            "created_at": self.created_at,
            "last_updated": self.last_updated,
            "turn_count": self.turn_count,
            "full_name": self.full_name,
            "full_name_confidence": self.full_name_confidence,
            "email": self.email,
            "email_confidence": self.email_confidence,
            "phone": self.phone,
            "phone_confidence": self.phone_confidence,
            "job_title": self.job_title,
            "job_title_confidence": self.job_title_confidence,
            "experience_level": self.experience_level,
            "experience_level_confidence": self.experience_level_confidence,
            "summary": self.summary,
            "summary_confidence": self.summary_confidence,
            "experience": self.experience,
            "education": self.education,
            "skills": self.skills,
            "projects": self.projects,
            "competencies": self.competencies,
            "community": self.community,
            "certifications": self.certifications,
            "fields_filled": list(self.fields_filled),
            "fields_missing": list(self.fields_missing),
            "progress_percentage": self.calculate_progress()
        }
    
    def calculate_progress(self) -> int:
        """Calculate completion percentage."""
        required_fields = {"full_name", "job_title", "experience", "education", "skills"}
        if not required_fields:
            return 0
        filled = len(required_fields & self.fields_filled)
        return int((filled / len(required_fields)) * 100)
    
    def to_resume_dict(self) -> Dict:
        """Convert to the resume format used by the builder."""
        return {
            "full_name": self.full_name,
            "email": self.email,
            "phone": self.phone,
            "job_title": self.job_title,
            "experience_level": self.experience_level,
            "summary": self.summary,
            "experience": self.experience,
            "education": self.education,
            "skills": self.skills,
            "projects": self.projects,
            "competencies": self.competencies,
            "community": self.community,
            "certifications": self.certifications
        }


def get_or_create_resume_state(session_id: str = "") -> ResumeState:
    """Get existing session or create new one."""
    if session_id and session_id in resume_sessions:
        return resume_sessions[session_id]
    
    state = ResumeState(session_id)
    resume_sessions[state.session_id] = state
    return state


# Groq Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")


def build_context_prompt(state: ResumeState, transcript: str) -> str:
    """Build dynamic prompt based on what we know and what's missing."""
    
    # Build state description
    state_desc = f"""CURRENT RESUME STATE:
- full_name: {state.full_name or 'MISSING'} {'✓' if state.full_name else '✗'}
- email: {state.email or 'MISSING'} {'✓' if state.email else '✗'}
- phone: {state.phone or 'MISSING'} {'✓' if state.phone else '✗'}
- job_title: {state.job_title or 'MISSING'} {'✓' if state.job_title else '✗'}
- experience_level: {state.experience_level} {'✓' if state.experience_level != 'entry' or state.experience else '✗'}
- summary: {'✓ PRESENT' if state.summary else '✗ MISSING'}
- experience: {len(state.experience)} job(s) {'✓' if state.experience else '✗'}
- education: {len(state.education)} school(s) {'✓' if state.education else '✗'}
- skills: {len(state.skills)} skill(s) {'✓' if state.skills else '✗'}

MISSING FIELDS TO FOCUS ON: {', '.join(state.fields_missing) if state.fields_missing else 'NONE - resume is complete!'}
"""

    # Build focus instruction based on what's missing
    focus_instruction = ""
    if not state.full_name:
        focus_instruction += "PRIORITY 1: Extract the person's FULL NAME from this transcript.\n"
    if not state.email:
        focus_instruction += "PRIORITY 2: Extract EMAIL ADDRESS. Look for patterns like user@domain.com.\n"
    if not state.phone:
        focus_instruction += "PRIORITY 3: Extract PHONE NUMBER. Any format works.\n"
    if not state.job_title:
        focus_instruction += "PRIORITY 4: Determine the person's TARGET JOB TITLE or current role.\n"
    if not state.experience:
        focus_instruction += "PRIORITY 5: Extract WORK EXPERIENCE - jobs, companies, dates, descriptions.\n"
    if state.experience and not state.education:
        focus_instruction += "PRIORITY 6: Extract EDUCATION - schools, degrees, fields, dates.\n"
    if not state.skills:
        focus_instruction += "PRIORITY 7: Extract SKILLS - technical, soft skills, tools, certifications.\n"
    
    # Handle promotions and special cases
    special_rules = """
SPECIAL RULES:
- PROMOTIONS: If user says "promoted to X", create a NEW job entry with new title.
  Set previous job end date to when promotion happened.
  Example: "Cashier 2020-2023, then promoted to Supervisor" → 
    Job 1: Cashier at Walmart, 2020-2023
    Job 2: Supervisor at Walmart, 2023-Present

- DATE INFERENCE:
  - "3 years" with no start → estimate from context
  - "since 2020" → 2020-Present
  - "promoted after 2 years" → new job starts 2 years after old job start

- CURRENT JOB: If user uses present tense or no end date, assume current job.

- SUMMARY: If missing, generate from what we know after extracting new info.
"""

    return f"""You are an expert resume builder AI. You are helping a user build their resume by extracting information from their voice transcript.

{state_desc}

{focus_instruction}

{special_rules}

OUTPUT REQUIREMENTS:
- Output ONLY valid JSON
- Use EXACTLY these field names: full_name, email, phone, job_title, experience_level, summary, experience (array), education (array), skills (array)
- Each job in experience must have: title, company, dates, description
- Each school in education must have: school, degree, field, dates
- If a field is already filled (✓ above), include its current value in output
- If a field is missing (✗ above), extract it from the transcript
- If you can't extract something, use empty string ""
- DO NOT use markdown code blocks
- DO NOT add explanations

The user is about to speak. Extract what you can and return the complete resume state.

User transcript:
{transcript}"""


def merge_parsed_into_state(state: ResumeState, parsed: Dict) -> ResumeState:
    """Smart merge: update state with new parsed data, preserving existing info."""
    
    # Update scalar fields (only if new data exists and old is empty)
    for field in ["full_name", "email", "phone", "job_title", "summary"]:
        new_val = parsed.get(field, "")
        old_val = getattr(state, field)
        if new_val and not old_val:
            state.update_field(field, new_val, confidence=0.85)
        elif new_val and old_val and new_val != old_val:
            # User provided update - increase confidence
            state.update_field(field, new_val, confidence=0.95)
    
    # Experience: merge jobs intelligently
    new_jobs = parsed.get("experience", [])
    for job in new_jobs:
        if job.get("company") and job.get("title"):
            state.add_experience(job)
    
    # Education: merge schools
    new_edu = parsed.get("education", [])
    for edu in new_edu:
        if edu.get("school"):
            state.add_education(edu)
    
    # Skills: add and deduplicate
    new_skills = parsed.get("skills", [])
    if new_skills:
        state.add_skills(new_skills)
    
    # Experience level: infer if we have jobs
    if state.experience and not state.experience_level_confidence:
        state.infer_experience_level()
    
    # Turn count increment
    state.turn_count += 1
    
    return state


def build_parser_prompt(transcript: str) -> str:
    """DEPRECATED: Use build_context_prompt instead. Kept for backward compatibility."""
    return build_context_prompt(ResumeState(), transcript)


async def parse_with_groq(transcript: str, state: ResumeState = None) -> dict:
    """Send transcript to Groq API and return parsed resume data.
    
    If state is provided, uses context-aware prompting.
    """
    if not GROQ_API_KEY:
        print("[GROQ] No API key configured")
        return {}
    
    # Use context-aware prompt if we have state
    if state:
        prompt = build_context_prompt(state, transcript)
        system_msg = "You are an expert resume builder AI that tracks state and extracts information intelligently."
    else:
        prompt = build_parser_prompt(transcript)
        system_msg = "You are a resume parser. Extract structured information from natural language. Output valid JSON only."
    
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 2000,
                "response_format": {"type": "json_object"}
            }
        )
        
        if response.status_code != 200:
            print(f"[GROQ] Error: {response.status_code} - {response.text}")
            return {}
        
        data = response.json()
        content = data["choices"][0]["message"]["content"]
        
        try:
            parsed = json.loads(content)
            return parsed
        except json.JSONDecodeError:
            print(f"[GROQ] Invalid JSON response: {content}")
            return {}


@app.post("/api/parse-voice")
async def parse_voice(request: Request):
    """Parse voice transcript into structured resume data using context-aware prompting."""
    try:
        data = await request.json()
        transcript = data.get("transcript", "").strip()
        session_id = data.get("session_id", "")
        
        if not transcript:
            return JSONResponse({"error": "Transcript is required"}, status_code=400)
        
        if len(transcript) < 20:
            return JSONResponse({"error": "Please provide more details (at least 20 characters)"}, status_code=400)
        
        # Get or create session state
        state = get_or_create_resume_state(session_id)
        
        # Parse with Groq using context-aware prompt
        parsed = await parse_with_groq(transcript, state=state)
        
        if not parsed:
            return JSONResponse(
                {"error": "Failed to parse resume. Please try again or provide more details."},
                status_code=500
            )
        
        # Smart merge into state
        state = merge_parsed_into_state(state, parsed)
        
        # Check if we should generate summary
        if not state.summary and (state.experience or state.education):
            state.summary = await generate_summary_from_state(state)
            if state.summary:
                state.fields_filled.add("summary")
                state.fields_missing.discard("summary")
        
        # Store in traditional resume format for backward compatibility
        resume_data = state.to_resume_dict()
        
        # Generate a resume ID (for the builder page)
        resume_id = secrets.token_urlsafe(16)
        resumes[resume_id] = {
            "id": resume_id,
            "data": resume_data,
            "created_at": datetime.now().isoformat(),
            "source": "voice",
            "session_id": state.session_id,
            "paid": False
        }
        
        return {
            "success": True,
            "resume_id": resume_id,
            "session_id": state.session_id,
            "data": resume_data,
            "progress": state.calculate_progress(),
            "fields_filled": list(state.fields_filled),
            "fields_missing": list(state.fields_missing),
            "turn_count": state.turn_count
        }
        
    except Exception as e:
        print(f"[PARSE VOICE] Error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)


async def generate_summary_from_state(state: ResumeState) -> str:
    """Generate a professional summary from the current resume state."""
    if not GROQ_API_KEY:
        return ""
    
    # Build context for summary generation
    jobs_text = ""
    for job in state.experience[:3]:  # Top 3 jobs
        jobs_text += f"- {job.get('title', '')} at {job.get('company', '')}\n"
    
    edu_text = ""
    for edu in state.education[:2]:
        edu_text += f"- {edu.get('degree', '')} in {edu.get('field', '')} from {edu.get('school', '')}\n"
    
    skills_text = ", ".join(state.skills[:10])
    
    prompt = f"""Write a 2-3 sentence professional summary for this person.

Name: {state.full_name}
Target Role: {state.job_title}
Experience: {len(state.experience)} job(s)
Top Skills: {skills_text}

Recent Work:
{jobs_text}

Education:
{edu_text}

Write in third person, professional tone. Highlight key strengths and career trajectory.
Output ONLY the summary text."""
    
    async with httpx.AsyncClient(timeout=15.0) as client:
        response = await client.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 500
            }
        )
        
        if response.status_code == 200:
            data = response.json()
            return data["choices"][0]["message"]["content"].strip()
    
    return ""


@app.get("/api/resume-state/{session_id}")
async def get_resume_state(session_id: str):
    """Get current resume state for a session."""
    if session_id in resume_sessions:
        state = resume_sessions[session_id]
        return {
            "success": True,
            "data": state.to_dict()
        }
    return JSONResponse({"error": "Session not found"}, status_code=404)


@app.get("/voice", response_class=HTMLResponse)
async def voice_page(request: Request):
    """Voice-first resume builder page"""
    return templates.TemplateResponse(request=request, name="voice_chat.html")


# === Voice feature ends ===


@app.get("/api/download/{resume_id}")
async def download_resume(resume_id: str, format: str = "docx"):
    """Download resume as .docx or .pdf"""
    output_dir = os.path.join(tempfile.gettempdir(), "resumes")
    
    if format.lower() == "pdf":
        filepath = os.path.join(output_dir, f"{resume_id}.pdf")
        filename = f"resume_{resume_id}.pdf"
    else:
        filepath = os.path.join(output_dir, f"{resume_id}.docx")
        filename = f"resume_{resume_id}.docx"
    
    if os.path.exists(filepath):
        return FileResponse(filepath, filename=filename)
    return JSONResponse({"error": "Resume not found"}, status_code=404)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
